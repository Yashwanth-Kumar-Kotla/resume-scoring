"""
Parsers for resume and job descriptions
"""

from docx import Document
from typing import Dict, List, Any
import re
import json


class ResumeParser:
    """Parse .docx resume files"""

    @staticmethod
    def parse_docx(file_path: str) -> Dict[str, Any]:
        """Parse resume from .docx file"""
        doc = Document(file_path)

        text = "\n".join([para.text for para in doc.paragraphs])

        return {
            "text": text,
            "doc": doc,
            "paragraphs": [para.text for para in doc.paragraphs],
            "word_count": len(text.split()),
            "page_count": len(doc.paragraphs) // 20,  # Rough estimate
        }

    @staticmethod
    def extract_sections(resume_text: str) -> Dict[str, str]:
        """Extract different sections from resume"""
        sections = {}

        # Common section headers
        section_headers = {
            "contact": r"(?:contact|phone|email)",
            "summary": r"(?:summary|objective|professional summary)",
            "experience": r"(?:experience|employment|work|professional experience)",
            "education": r"(?:education|degree|university|college)",
            "skills": r"(?:skills|technical skills|competencies)",
            "projects": r"(?:projects|portfolio)",
            "publications": r"(?:publications|papers|articles)",
        }

        text_lower = resume_text.lower()

        for section_name, pattern in section_headers.items():
            if re.search(pattern, text_lower):
                sections[section_name] = "present"

        return sections


class JDParser:
    """Parse job descriptions and extract skills"""

    def __init__(self):
        self.common_skills = {
            # Languages
            "python", "java", "javascript", "c++", "go", "rust", "typescript",
            "sql", "r", "scala", "kotlin", "swift", "objective-c",

            # Frameworks & Libraries
            "react", "angular", "vue", "django", "flask", "spring", "fastapi",
            "node.js", "express", "asp.net", "rails", "laravel", "symfony",
            "tensorflow", "pytorch", "scikit-learn", "keras", "pandas", "numpy",

            # Databases
            "postgresql", "mysql", "mongodb", "redis", "dynamodb", "cassandra",
            "oracle", "mssql", "elasticsearch",

            # Cloud & DevOps
            "aws", "gcp", "azure", "docker", "kubernetes", "jenkins", "gitlab-ci",
            "github actions", "terraform", "ansible", "lambda", "ec2", "s3",

            # Big Data & ML
            "spark", "hadoop", "kafka", "machine learning", "deep learning",
            "nlp", "computer vision", "data science", "analytics",

            # Other
            "git", "agile", "scrum", "rest api", "microservices", "ci/cd",
            "linux", "unix", "windows", "unix", "etl", "data warehouse",
            "oops", "design patterns", "solid principles"
        }

    def extract_skills(self, jd_text: str) -> List[str]:
        """Extract skills from job description"""
        skills_found = []
        jd_lower = jd_text.lower()

        for skill in self.common_skills:
            if skill in jd_lower:
                skills_found.append(skill)

        return list(set(skills_found))  # Remove duplicates

    def categorize_skills(self, skills: List[str]) -> Dict[str, List[str]]:
        """Categorize skills by type"""
        categories = {
            "languages": [],
            "frameworks": [],
            "databases": [],
            "cloud": [],
            "other": []
        }

        skill_categories = {
            "languages": ["python", "java", "javascript", "c++", "go", "rust", "typescript",
                         "sql", "r", "scala", "kotlin", "swift", "objective-c"],
            "frameworks": ["react", "angular", "vue", "django", "flask", "spring", "fastapi",
                          "node.js", "express", "tensorflow", "pytorch", "scikit-learn"],
            "databases": ["postgresql", "mysql", "mongodb", "redis", "dynamodb", "cassandra"],
            "cloud": ["aws", "gcp", "azure", "docker", "kubernetes", "jenkins", "ci/cd"],
        }

        for skill in skills:
            categorized = False
            for category, skill_list in skill_categories.items():
                if skill in skill_list:
                    categories[category].append(skill)
                    categorized = True
                    break

            if not categorized:
                categories["other"].append(skill)

        return categories

    def assess_relevance(self, skills: List[str], resume_text: str) -> Dict[str, List[str]]:
        """Assess which skills from JD are likely already in resume"""
        resume_lower = resume_text.lower()

        present = []
        missing = []

        for skill in skills:
            if skill in resume_lower:
                present.append(skill)
            else:
                missing.append(skill)

        return {
            "present": present,
            "missing": missing
        }

    def prioritize_skills(self, jd_text: str, skills: List[str]) -> List[str]:
        """Prioritize skills by frequency in JD"""
        jd_lower = jd_text.lower()

        skill_frequency = {}
        for skill in skills:
            # Count occurrences
            count = jd_lower.count(skill)
            skill_frequency[skill] = count

        # Sort by frequency
        sorted_skills = sorted(skill_frequency.items(), key=lambda x: x[1], reverse=True)
        return [skill for skill, _ in sorted_skills]


class SkillVerifier:
    """Help user verify skills they have"""

    def __init__(self):
        self.skill_contexts = {
            "python": ["data science", "backend", "automation", "ml"],
            "java": ["backend", "enterprise", "android"],
            "javascript": ["frontend", "react", "node"],
            "sql": ["database", "data", "analytics"],
            "aws": ["cloud", "infrastructure", "devops"],
            "docker": ["containerization", "devops", "infrastructure"],
            "machine learning": ["data science", "ai", "ml"],
            "react": ["frontend", "web development"],
            "kubernetes": ["infrastructure", "devops", "orchestration"],
            "oops": ["programming", "design", "foundations"],
        }

    def create_verification_form(self, missing_skills: List[str]) -> Dict[str, Dict]:
        """Create verification form for user"""
        form = {}

        for skill in missing_skills:
            form[skill] = {
                "skill": skill,
                "context": self.skill_contexts.get(skill, ["general"]),
                "verified": False,
                "level": None,  # beginner, intermediate, advanced
                "source": None,  # university, personal project, internship, etc.
            }

        return form

    def process_verification(self, verification_data: Dict) -> Dict:
        """Process user verification responses"""
        verified_skills = {}
        unverified_skills = {}

        for skill, data in verification_data.items():
            if data.get("verified"):
                verified_skills[skill] = {
                    "level": data.get("level"),
                    "source": data.get("source"),
                }
            else:
                unverified_skills[skill] = data

        return {
            "verified": verified_skills,
            "unverified": unverified_skills,
        }
