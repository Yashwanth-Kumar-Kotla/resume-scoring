"""
Writers for output files (resume, reports)
"""

from docx import Document
from docx.shared import Pt, Inches
from typing import Dict, List, Any
import json
from datetime import datetime


class ResumeWriter:
    """Write modified resume back to .docx preserving formatting"""

    @staticmethod
    def update_resume(original_doc_path: str, modifications: Dict, output_path: str) -> str:
        """Update resume with modifications while preserving formatting"""
        doc = Document(original_doc_path)

        # Track changes for reporting
        changes = []

        # Apply modifications to document
        if "updated_text" in modifications:
            updated_text = modifications["updated_text"]
            para_updates = modifications.get("paragraph_updates", {})

            # Update specific paragraphs if provided
            for para_index, new_text in para_updates.items():
                try:
                    para_index = int(para_index)
                    if para_index < len(doc.paragraphs):
                        old_text = doc.paragraphs[para_index].text
                        doc.paragraphs[para_index].text = new_text
                        changes.append({
                            "type": "text_update",
                            "paragraph": para_index,
                            "old": old_text,
                            "new": new_text
                        })
                except (ValueError, IndexError):
                    pass

        # Save document
        doc.save(output_path)

        return output_path

    @staticmethod
    def preserve_formatting(source_doc, modified_text: str) -> Document:
        """Create new doc with modified text while preserving original formatting"""
        doc = Document(source_doc)

        # Keep original formatting but update text
        # This is a simplified approach - keeps structure, updates content
        return doc


class ReportWriter:
    """Generate comprehensive reports"""

    @staticmethod
    def generate_ats_report(ats_results: Dict, output_path: str) -> str:
        """Generate ATS scoring report"""
        content = []

        content.append("=" * 80)
        content.append("📊 ATS SCORING REPORT")
        content.append("=" * 80)
        content.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        content.append("")

        # Calculate average
        scores = [r.overall_score for r in ats_results.values()]
        avg_score = sum(scores) / len(scores) if scores else 0

        content.append(f"🎯 AVERAGE ATS SCORE: {avg_score:.0f}/100")
        content.append(f"Systems Passed: {sum(1 for r in ats_results.values() if r.overall_score >= 80)}/6")
        content.append("")

        # Detailed scores
        content.append("📈 DETAILED SCORES BY ATS SYSTEM:")
        content.append("-" * 80)

        for system_name, score in sorted(ats_results.items(), key=lambda x: x[1].overall_score, reverse=True):
            content.append(f"\n{system_name}: {score.overall_score}/100 - {score.verdict}")
            content.append(f"  Formatting: {score.formatting}/100")
            content.append(f"  Keywords:   {score.keywords}/100")
            content.append(f"  Sections:   {score.sections}/100")
            content.append(f"  Experience: {score.experience}/100")
            content.append(f"  Education:  {score.education}/100")

            if score.suggestions:
                content.append(f"  Suggestions:")
                for suggestion in score.suggestions:
                    content.append(f"    • {suggestion}")

        # Write to file
        with open(output_path, 'w') as f:
            f.write("\n".join(content))

        return output_path

    @staticmethod
    def generate_holistic_report(holistic_score: Any, output_path: str) -> str:
        """Generate holistic evaluation report"""
        content = []

        content.append("=" * 80)
        content.append("📊 HOLISTIC RESUME EVALUATION")
        content.append("=" * 80)
        content.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        content.append("")

        # Overall score
        content.append(f"🎯 OVERALL SCORE: {holistic_score.overall_score:.1f}/{holistic_score.max_score}")
        content.append("")

        # Detailed scores
        content.append("📈 DETAILED SCORES:")
        content.append("-" * 80)
        content.append(f"🌐 Open Source: {holistic_score.open_source.score:.1f}/{holistic_score.open_source.max_score}")
        content.append(f"   Evidence: {holistic_score.open_source.evidence}")
        content.append("")

        content.append(f"🚀 Self Projects: {holistic_score.self_projects.score:.1f}/{holistic_score.self_projects.max_score}")
        content.append(f"   Evidence: {holistic_score.self_projects.evidence}")
        content.append("")

        content.append(f"🏢 Production Experience: {holistic_score.production_experience.score:.1f}/{holistic_score.production_experience.max_score}")
        content.append(f"   Evidence: {holistic_score.production_experience.evidence}")
        content.append("")

        content.append(f"💻 Technical Skills: {holistic_score.technical_skills.score:.1f}/{holistic_score.technical_skills.max_score}")
        content.append(f"   Evidence: {holistic_score.technical_skills.evidence}")
        content.append("")

        # Bonuses and deductions
        if holistic_score.bonus_points > 0:
            content.append(f"⭐ BONUS POINTS: {holistic_score.bonus_points:.1f}")
            content.append("")

        if holistic_score.deductions > 0:
            content.append(f"⚠️ DEDUCTIONS: -{holistic_score.deductions:.1f}")
            content.append("")

        # Strengths
        if holistic_score.key_strengths:
            content.append("✅ KEY STRENGTHS:")
            for i, strength in enumerate(holistic_score.key_strengths, 1):
                content.append(f"  {i}. {strength}")
            content.append("")

        # Areas for improvement
        if holistic_score.areas_for_improvement:
            content.append("🔧 AREAS FOR IMPROVEMENT:")
            for i, area in enumerate(holistic_score.areas_for_improvement, 1):
                content.append(f"  {i}. {area}")

        # Write to file
        with open(output_path, 'w') as f:
            f.write("\n".join(content))

        return output_path

    @staticmethod
    def generate_comprehensive_report(ats_results: Dict, holistic_score: Any,
                                     skill_analysis: Dict, output_path: str) -> str:
        """Generate comprehensive report combining all analyses"""
        content = []

        content.append("=" * 100)
        content.append(" " * 30 + "🎯 COMPREHENSIVE RESUME ANALYSIS REPORT")
        content.append("=" * 100)
        content.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        content.append("")

        # SECTION 1: ATS Analysis
        content.append("\n" + "=" * 100)
        content.append("SECTION 1: ATS SYSTEM COMPATIBILITY")
        content.append("=" * 100)

        avg_ats = sum(r.overall_score for r in ats_results.values()) / len(ats_results)
        content.append(f"\n📊 Average ATS Score: {avg_ats:.0f}/100")
        content.append(f"Systems Passed: {sum(1 for r in ats_results.values() if r.overall_score >= 80)}/6")
        content.append("")

        for system_name, score in sorted(ats_results.items(), key=lambda x: x[1].overall_score, reverse=True):
            status = "✅" if score.overall_score >= 80 else "⚠️" if score.overall_score >= 60 else "❌"
            content.append(f"{status} {system_name}: {score.overall_score}/100")

        # SECTION 2: Holistic Evaluation
        content.append("\n\n" + "=" * 100)
        content.append("SECTION 2: HOLISTIC EVALUATION")
        content.append("=" * 100)

        content.append(f"\n🎯 Overall Score: {holistic_score.overall_score:.1f}/{holistic_score.max_score}")
        content.append("")
        content.append("Category Breakdown:")
        content.append(f"  • Open Source: {holistic_score.open_source.score:.1f}/{holistic_score.open_source.max_score}")
        content.append(f"  • Self Projects: {holistic_score.self_projects.score:.1f}/{holistic_score.self_projects.max_score}")
        content.append(f"  • Production Experience: {holistic_score.production_experience.score:.1f}/{holistic_score.production_experience.max_score}")
        content.append(f"  • Technical Skills: {holistic_score.technical_skills.score:.1f}/{holistic_score.technical_skills.max_score}")

        # SECTION 3: Skills Analysis
        content.append("\n\n" + "=" * 100)
        content.append("SECTION 3: SKILLS ANALYSIS")
        content.append("=" * 100)

        if "missing_skills" in skill_analysis:
            content.append(f"\n📋 Missing Skills from Job Description: {len(skill_analysis['missing_skills'])}")
            for i, skill in enumerate(skill_analysis['missing_skills'][:10], 1):
                content.append(f"  {i}. {skill}")

        if "present_skills" in skill_analysis:
            content.append(f"\n✅ Already Present Skills: {len(skill_analysis['present_skills'])}")
            for i, skill in enumerate(skill_analysis['present_skills'][:10], 1):
                content.append(f"  {i}. {skill}")

        # SECTION 4: Recommendations
        content.append("\n\n" + "=" * 100)
        content.append("SECTION 4: RECOMMENDATIONS")
        content.append("=" * 100)

        if holistic_score.areas_for_improvement:
            content.append("\n🎯 Priority Focus Areas:")
            for i, area in enumerate(holistic_score.areas_for_improvement, 1):
                content.append(f"  {i}. {area}")

        content.append("\n💡 Next Steps:")
        content.append("  1. Verify which missing skills you actually have")
        content.append("  2. Add relevant keywords to your resume")
        content.append("  3. Strengthen bullet points with quantifiable metrics")
        content.append("  4. Ensure resume stays within 1 page")
        content.append("  5. Re-scan with updated resume to verify improvements")

        # Write to file
        with open(output_path, 'w') as f:
            f.write("\n".join(content))

        return output_path
