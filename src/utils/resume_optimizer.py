"""
Resume Optimization Module - Rewrites and enhances resume with keywords
"""

import anthropic
import os
from typing import Dict, List, Any
from docx import Document
from docx.shared import Pt


class ResumeOptimizer:
    """Optimize resume by adding keywords and strengthening bullets"""

    def __init__(self):
        self.client = None
        api_key = os.getenv("ANTHROPIC_API_KEY")
        if api_key and api_key != "your_api_key_here":
            try:
                self.client = anthropic.Anthropic(api_key=api_key)
            except:
                pass

    def optimize_resume(self, resume_text: str, verified_skills: Dict,
                       jd_text: str = None) -> str:
        """
        Optimize resume by adding verified skills and strengthening bullets
        Returns optimized resume text
        """

        verified_skill_list = [skill for skill, data in verified_skills.items()
                             if data.get('verified')]

        if not verified_skill_list:
            return resume_text

        # Use Claude if available, otherwise use basic optimization
        if self.client:
            return self._optimize_with_claude(resume_text, verified_skill_list, jd_text)
        else:
            return self._optimize_basic(resume_text, verified_skill_list)

    def _optimize_with_claude(self, resume_text: str, verified_skills: List[str],
                             jd_text: str) -> str:
        """Use Claude to intelligently rewrite resume"""

        skills_str = ", ".join(verified_skills)

        prompt = f"""You are a professional resume writer. Improve this resume by:

1. Adding these verified skills intelligently to appropriate sections: {skills_str}
2. Strengthening bullet points with action verbs and metrics
3. Maintaining professional tone and formatting
4. Keeping to approximately 1 page (max 500 words)
5. Adding quantifiable impact where possible
6. NOT making up experience the person doesn't have

ORIGINAL RESUME:
{resume_text}

{'JOB DESCRIPTION (for context):' if jd_text else ''}
{jd_text[:1000] if jd_text else ''}

IMPORTANT:
- Only add the skills: {skills_str}
- Place each skill in the most relevant section
- Don't add skills to unrelated sections
- Strengthen existing bullets, don't remove them
- Keep the overall structure and content truthful
- Ensure it still fits on 1 page

Return ONLY the improved resume text, no explanations."""

        try:
            message = self.client.messages.create(
                model="claude-opus-4-8",
                max_tokens=2000,
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
            optimized = message.content[0].text
            return optimized
        except Exception as e:
            print(f"⚠️ Claude optimization failed: {e}")
            return self._optimize_basic(resume_text, verified_skills)

    def _optimize_basic(self, resume_text: str, verified_skills: List[str]) -> str:
        """Basic optimization without Claude"""

        optimized = resume_text

        # Add skills to skills section
        for skill in verified_skills:
            if skill.lower() not in optimized.lower():
                # Try to add to skills section
                if "skill" in optimized.lower():
                    # Simple insertion - find skills section and add
                    lines = optimized.split("\n")
                    for i, line in enumerate(lines):
                        if "skill" in line.lower() and i + 1 < len(lines):
                            # Add skill to next line
                            if lines[i + 1].strip():
                                lines[i + 1] = lines[i + 1].rstrip() + f", {skill}"
                            else:
                                lines[i + 1] = f"  {skill}"
                            optimized = "\n".join(lines)
                            break

        return optimized

    def save_optimized_resume(self, original_doc_path: str, optimized_text: str,
                             output_path: str, verified_skills: List[str] = None) -> Dict:
        """
        Smart keyword insertion: Add to Skills section when they don't fit experience bullets.
        Only add to bullets if they genuinely describe that work.
        Returns: {success, output_path, keywords_added, keywords_added_location, verified, verification_proof}
        """
        try:
            doc = Document(original_doc_path)
            result = {
                "success": False,
                "output_path": output_path,
                "keywords_added": [],
                "keywords_added_location": {},  # Track where each keyword was added
                "verified": [],
                "verification_proof": []
            }

            if not verified_skills:
                verified_skills = []

            # Step 1: Find the Skills section
            skills_section_index = None
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip().lower() in ["skills", "technical skills", "core skills"]:
                    skills_section_index = i
                    break

            # Step 2: Map verified skills to where they should go
            # RULE: Only add to experience bullets if the keyword genuinely describes that specific work
            skills_for_bullets = {
                # keyword: (paragraph_contains, replacement_rule)
                "docker": ("ci/cd", lambda t: t.replace("pipeline (test", "pipeline with Docker containers (test")),
                "threshold optimization": ("calibration", lambda t: t.replace("decision threshold", "decision threshold optimization")),
                "smote": ("imbalanced", lambda t: t),  # Already there, don't force
            }

            skills_for_skills_section = set(verified_skills) - set(skills_for_bullets.keys())

            # Step 3: Add to experience bullets only if it makes sense
            for para in doc.paragraphs:
                para_text = para.text.lower()

                for skill, (contains_keyword, replacement_fn) in skills_for_bullets.items():
                    if contains_keyword in para_text and skill not in para_text:
                        # Try the replacement
                        for run in para.runs:
                            if contains_keyword in run.text.lower():
                                original = run.text
                                run.text = replacement_fn(run.text)
                                if run.text != original:  # Only mark if changed
                                    result["keywords_added"].append(skill)
                                    result["keywords_added_location"][skill] = "experience bullet"
                                break

            # Step 4: Add remaining skills to Skills section (SAFE PLACE for keywords)
            if skills_section_index is not None and skills_for_skills_section:
                skills_para = doc.paragraphs[skills_section_index]

                # Find and update each skill subsection
                for skill in skills_for_skills_section:
                    added = False

                    # Try to add to most relevant category
                    if skill.lower() in ["python", "sql", "javascript", "java", "go", "rust"]:
                        category = "Languages:"
                    elif skill.lower() in ["xgboost", "pytorch", "tensorflow", "scikit-learn", "deep learning", "computer vision", "nlp"]:
                        category = "ML & Statistics:"
                    elif skill.lower() in ["aws", "gcp", "azure", "kubernetes", "docker"]:
                        category = "Cloud:"
                    elif skill.lower() in ["langchain", "rag", "prompt engineering"]:
                        category = "LLM & GenAI:"
                    else:
                        category = "Data Engineering & Deployment:"

                    # Look for the category line after skills section
                    for i in range(skills_section_index + 1, len(doc.paragraphs)):
                        para = doc.paragraphs[i]
                        if category.lower() in para.text.lower():
                            # Add skill to this line
                            for run in para.runs:
                                if category.lower() in run.text.lower():
                                    if skill not in run.text:
                                        run.text = run.text.rstrip() + f", {skill}"
                                        result["keywords_added"].append(skill)
                                        result["keywords_added_location"][skill] = "skills section"
                                        added = True
                                    break
                            if added:
                                break

            # Step 5: SAVE
            doc.save(output_path)

            # Step 6: IMMEDIATE VERIFICATION
            verify_doc = Document(output_path)
            verify_text = "\n".join([p.text for p in verify_doc.paragraphs])

            # Check each keyword
            for kw in result["keywords_added"]:
                if kw.lower() in verify_text.lower():
                    result["verified"].append(kw)
                    location = result["keywords_added_location"].get(kw, "unknown")
                    # Find proof
                    for para in verify_doc.paragraphs:
                        if kw.lower() in para.text.lower():
                            result["verification_proof"].append({
                                "keyword": kw,
                                "location": location,
                                "found_in": para.text[:100]
                            })
                            break

            # Success only if all keywords verified
            result["success"] = len(result["verified"]) == len(result["keywords_added"])

            return result

        except Exception as e:
            print(f"Error saving optimized resume: {e}")
            return {
                "success": False,
                "output_path": output_path,
                "keywords_added": [],
                "keywords_added_location": {},
                "verified": [],
                "verification_proof": [],
                "error": str(e)
            }


class ResumeTailorizer:
    """Tailor resume for specific job"""

    def __init__(self):
        self.optimizer = ResumeOptimizer()

    def tailor_for_job(self, resume_path: str, jd_text: str,
                      verified_skills: Dict, output_dir: str = "output") -> Dict:
        """
        Tailor resume for specific job with verification.
        Returns: {success, optimized_path, keywords_added, verified_in_file, verification_proof}
        """

        from .parsers import ResumeParser

        # Parse original resume
        resume_data = ResumeParser.parse_docx(resume_path)
        resume_text = resume_data["text"]

        # Optimize resume with Claude
        optimized_text = self.optimizer.optimize_resume(
            resume_text,
            verified_skills,
            jd_text
        )

        # Save optimized resume with immediate verification
        optimized_path = f"{output_dir}/resume_tailored.docx"
        verified_skills_list = [s for s, d in verified_skills.items() if d.get('verified')]

        save_result = self.optimizer.save_optimized_resume(
            resume_path,
            optimized_text,
            optimized_path,
            verified_skills_list
        )

        # Save text versions for reference
        with open(f"{output_dir}/resume_original.txt", 'w') as f:
            f.write(resume_text)

        with open(f"{output_dir}/resume_optimized.txt", 'w') as f:
            f.write(optimized_text)

        # Return with verification proof
        return {
            "success": save_result["success"],
            "optimized_path": optimized_path,
            "keywords_added": save_result["keywords_added"],
            "verified_in_file": save_result["verified"],
            "verification_proof": save_result["verification_proof"],
            "all_keywords_confirmed": len(save_result["verified"]) == len(save_result["keywords_added"])
        }
